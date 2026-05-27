"""Load compliance source documents into text elements."""

from pathlib import Path
from typing import Iterable

from open_deep_research.compliance.schemas import ParsedDocumentElement, SourceType


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xlsm"}


def iter_source_files(root: str | Path) -> Iterable[Path]:
    """Yield files supported by the first compliance indexing pass."""

    root_path = Path(root)
    if not root_path.exists():
        return
    for path in root_path.rglob("*"):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield path


def load_file(path: str | Path, source_type: SourceType) -> list[ParsedDocumentElement]:
    """Load one supported file into traceable text elements."""

    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(file_path, source_type)
    if suffix == ".docx":
        return _load_docx(file_path, source_type)
    if suffix in {".xlsx", ".xlsm"}:
        return _load_xlsx(file_path, source_type)
    raise ValueError(f"Unsupported compliance document extension: {suffix}")


def _base_metadata(path: Path, source_type: SourceType) -> dict:
    return {
        "source_type": source_type,
        "path": str(path),
        "file_name": path.name,
        "extension": path.suffix.lower(),
    }


def _load_pdf(path: Path, source_type: SourceType) -> list[ParsedDocumentElement]:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("pymupdf is required to parse PDF files") from exc

    elements: list[ParsedDocumentElement] = []
    doc = fitz.open(path)
    try:
        for page_index, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            if not text:
                continue
            elements.append(
                ParsedDocumentElement(
                    **_base_metadata(path, source_type),
                    text=text,
                    page_number=page_index,
                )
            )
    finally:
        doc.close()
    return elements


def _load_docx(path: Path, source_type: SourceType) -> list[ParsedDocumentElement]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to parse DOCX files") from exc

    document = Document(path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                parts.append(row_text)

    if not parts:
        return []
    return [
        ParsedDocumentElement(
            **_base_metadata(path, source_type),
            text="\n".join(parts),
        )
    ]


def _load_xlsx(path: Path, source_type: SourceType) -> list[ParsedDocumentElement]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("openpyxl is required to parse XLSX/XLSM files") from exc

    workbook = load_workbook(path, read_only=True, data_only=True)
    elements: list[ParsedDocumentElement] = []
    try:
        for worksheet in workbook.worksheets:
            rows: list[str] = []
            first_row = None
            last_row = None
            for row in worksheet.iter_rows(values_only=True):
                values = [str(value).strip() for value in row if value is not None and str(value).strip()]
                if not values:
                    continue
                row_number = len(rows) + 1
                first_row = first_row or row_number
                last_row = row_number
                rows.append(" | ".join(values))
            if rows:
                elements.append(
                    ParsedDocumentElement(
                        **_base_metadata(path, source_type),
                        text="\n".join(rows),
                        sheet_name=worksheet.title,
                        row_range=f"{first_row}-{last_row}" if first_row and last_row else None,
                    )
                )
    finally:
        workbook.close()
    return elements

